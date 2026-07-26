"""Extract plain text from PDF and DOCX resume files."""

from __future__ import annotations

import io
import re
from pathlib import Path


class ExtractionError(Exception):
    """Raised when a resume file cannot be read."""


def extract_text(payload: bytes, file_name: str, mime_type: str | None = None) -> str:
    lower = file_name.lower()
    mime = (mime_type or "").lower()

    if lower.endswith(".pdf") or "pdf" in mime:
        return _extract_pdf(payload)
    if lower.endswith(".docx") or "wordprocessingml" in mime:
        return _extract_docx(payload)
    if lower.endswith(".doc") or mime == "application/msword":
        raise ExtractionError(
            "Legacy .doc files are not supported. Please upload PDF or DOCX."
        )

    raise ExtractionError("Only PDF and DOCX files are supported")


def _extract_pdf(payload: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("PDF support is not installed (pypdf)") from exc

    try:
        reader = PdfReader(io.BytesIO(payload))
        pages = list(reader.pages)
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Could not read PDF: {exc}") from exc

    # Layout mode keeps line and column structure, which every downstream heuristic
    # depends on. Many resumes otherwise extract as a single unbroken line.
    layout = _read_pages(pages, mode="layout")
    plain = _read_pages(pages, mode="plain")

    if _line_count(layout) > _line_count(plain):
        return _normalize(layout)

    return _normalize(plain)


def _read_pages(pages: list, mode: str) -> str:
    parts: list[str] = []
    for page in pages:
        try:
            if mode == "layout":
                parts.append(page.extract_text(extraction_mode="layout") or "")
            else:
                parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — a single bad page must not fail the upload
            continue
    return "\n".join(parts)


def _line_count(text: str) -> int:
    return sum(1 for line in text.splitlines() if line.strip())


def _extract_docx(payload: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("DOCX support is not installed (python-docx)") from exc

    try:
        document = Document(io.BytesIO(payload))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Could not read DOCX: {exc}") from exc

    return _normalize(text)


COLUMN_GAP = "   "


def _normalize(text: str) -> str:
    """Tidy whitespace while keeping wide gaps, which mark two-column layouts.

    Section headings in column layouts sit on the same line as their body
    ("SUMMARY      An experienced..."), so collapsing every run of spaces would
    destroy the only signal that separates them.
    """
    cleaned = text.replace("\x00", " ").replace("\t", " ")
    cleaned = re.sub(
        r" {2,}",
        lambda match: COLUMN_GAP if len(match.group(0)) >= 4 else " ",
        cleaned,
    )
    cleaned = re.sub(r"[ ]+\n", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def extension_for(file_name: str) -> str:
    return Path(file_name).suffix.lower()
